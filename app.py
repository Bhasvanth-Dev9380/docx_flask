import os
import requests
from io import BytesIO
from flask import Flask, request, send_file
import json
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

app = Flask(__name__)

# Function to replace the logo and company name in the document
def replace_logo_and_company_name(doc, logo_url, company_name):
    # Download the logo image from the URL
    response = requests.get(logo_url)
    image_bytes = BytesIO(response.content)

    # Access the header section of the document
    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            print(f"Paragraph before replacement: '{paragraph.text}'")  # Debugging: Print paragraph text

            # If both placeholders are in the same paragraph, we split and process them
            if "{logo}" in paragraph.text and "{Company}" in paragraph.text:
                # Split the paragraph text by {logo} and {Company}
                parts = paragraph.text.split("{logo}")
                before_logo = parts[0]  # Text before {logo}
                after_logo = parts[1].split("{Company}")  # Text after {logo} but before {Company}
                after_company = after_logo[1] if len(after_logo) > 1 else ""

                # Clear the paragraph and rebuild it
                paragraph.clear()

                # Add the text before logo
                run = paragraph.add_run(before_logo)

                # Insert the logo and ensure it doesn't cause layout issues
                run = paragraph.add_run()
                run.add_picture(image_bytes, width=Inches(0.35))

                # Add text between logo and company name
                paragraph.add_run(after_logo[0])

                # Add the company name and any remaining text after company name
                paragraph.add_run(company_name)
                paragraph.add_run(after_company)

                # Align paragraph to prevent any extra new lines or unwanted page breaks
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            print(f"Paragraph after replacement: '{paragraph.text}'")  # Debugging: Print paragraph after replacement


# Endpoint to handle document updating
@app.route('/update-document', methods=['POST'])
def update_document():
    # Ensure the uploads directory exists
    if not os.path.exists('uploads'):
        os.makedirs('uploads')

    # Get the JSON data from the form-data (it comes as a string, so we need to parse it)
    json_data = request.form['json']
    data = json.loads(json_data)  # Parse the JSON string

    # Get the uploaded document from the client
    uploaded_file = request.files['document']
    document_path = os.path.join('uploads', uploaded_file.filename)

    # Save the uploaded file locally
    uploaded_file.save(document_path)

    # Open and modify the Word document using python-docx
    doc = Document(document_path)

    # Extract company name and logo URL from JSON data
    company_name = data.get('company_name', 'Default Company')
    logo_url = data.get('logo_url', '')  # Assuming logo_url is a field in the incoming data

    # Replace the placeholders for the logo and company name
    if logo_url and company_name:
        replace_logo_and_company_name(doc, logo_url, company_name)

    entries = data.get('entries', [])

    # Assuming the document has a table
    for table in doc.tables:
        for entry in entries:
            found_row = False
            # Iterate over the rows and check for existing 'No.'
            for current_row_index in range(1, len(table.rows)):  # Skip header row
                row = table.rows[current_row_index]
                if row.cells[0].text.strip() == entry['No.']:  # Check 'No.' match
                    found_row = True
                    break

            if found_row:
                # 'No.' already exists, skip to next entry
                continue

            # If no match found, find the first empty row or add a new one
            for current_row_index in range(1, len(table.rows)):
                row = table.rows[current_row_index]
                if not row.cells[0].text.strip():  # Find empty row
                    break
            else:
                # If no empty row is found, add a new row
                last_row = table.rows[-1]
                row = table.add_row()
                for i, cell in enumerate(last_row.cells):
                    row.cells[i].text = cell.text  # Copy structure from the last row

            # Update the row with data from the JSON
            row.cells[0].text = entry['No.']  # No.
            row.cells[1].text = entry['Drawing Number']  # Drawing Number
            row.cells[2].text = entry['Drawing Title']  # Drawing Title
            row.cells[3].text = entry['Revision Number']  # Revision Number
            row.cells[4].text = entry['Date of Issue']  # Date of Issue
            row.cells[5].text = entry['Prepared By']  # Prepared By
            row.cells[6].text = entry['Approved By']  # Approved By
            row.cells[7].text = entry['Client Approval Status']  # Client Approval Status
            row.cells[8].text = entry['File Location/Reference']  # File Location/Reference
            row.cells[9].text = entry['Remarks']  # Remarks

    # Save the updated document
    updated_path = os.path.join('uploads', 'updated_' + uploaded_file.filename)
    doc.save(updated_path)

    # Send back the updated document
    return send_file(updated_path, as_attachment=True)



